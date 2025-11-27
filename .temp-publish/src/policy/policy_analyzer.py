"""
Policy Analyzer - Document Parsing and Policy Extraction

Purpose: Parse policy documents (PDF, Markdown, DOCX, TXT) and extract
         structured policy rules, requirements, and constraints.

Capabilities:
- Multi-format parsing (PDF, MD, DOCX, TXT)
- Policy rule extraction (MUST, SHOULD, MAY, MUST NOT, SHOULD NOT)
- Requirement categorization (security, quality, performance, architecture)
- Constraint identification (thresholds, limits, patterns)
- Metadata extraction (version, date, author, scope)

Author: Asif Hussain
Copyright: © 2024-2025 Asif Hussain. All rights reserved.
License: Source-Available (Use Allowed, No Contributions)
Repository: https://github.com/asifhussain60/CORTEX
"""

import os
import re
import hashlib
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
from dataclasses import dataclass, field, asdict
from enum import Enum


class PolicyLevel(Enum):
    """RFC 2119 policy requirement levels"""
    MUST = "MUST"
    MUST_NOT = "MUST NOT"
    SHOULD = "SHOULD"
    SHOULD_NOT = "SHOULD NOT"
    MAY = "MAY"
    RECOMMENDED = "RECOMMENDED"
    OPTIONAL = "OPTIONAL"


class PolicyCategory(Enum):
    """Policy categories for classification"""
    SECURITY = "security"
    QUALITY = "quality"
    PERFORMANCE = "performance"
    ARCHITECTURE = "architecture"
    TESTING = "testing"
    DOCUMENTATION = "documentation"
    DEPLOYMENT = "deployment"
    MAINTENANCE = "maintenance"
    COMPLIANCE = "compliance"
    GENERAL = "general"


@dataclass
class PolicyRule:
    """Structured policy rule"""
    id: str
    text: str
    level: PolicyLevel
    category: PolicyCategory
    keywords: List[str] = field(default_factory=list)
    threshold: Optional[float] = None
    unit: Optional[str] = None
    rationale: Optional[str] = None
    examples: List[str] = field(default_factory=list)
    source_line: Optional[int] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        result = asdict(self)
        result['level'] = self.level.value
        result['category'] = self.category.value
        return result


@dataclass
class PolicyDocument:
    """Parsed policy document"""
    file_path: str
    file_hash: str
    format: str  # pdf, md, docx, txt
    title: Optional[str] = None
    version: Optional[str] = None
    date: Optional[datetime] = None
    author: Optional[str] = None
    scope: Optional[str] = None
    rules: List[PolicyRule] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        result = asdict(self)
        result['rules'] = [rule.to_dict() for rule in self.rules]
        if self.date:
            result['date'] = self.date.isoformat()
        return result


class PolicyAnalyzer:
    """
    Analyze and parse policy documents into structured rules.
    
    Supports:
    - PDF parsing (PyPDF2, pdfplumber)
    - Markdown parsing (mistune)
    - DOCX parsing (python-docx)
    - TXT parsing (built-in)
    
    Extracts:
    - Policy rules with RFC 2119 levels (MUST, SHOULD, MAY, etc.)
    - Categories (security, quality, performance, etc.)
    - Thresholds and constraints
    - Metadata (version, date, author)
    """
    
    def __init__(self):
        """Initialize policy analyzer"""
        self.supported_formats = ['.pdf', '.md', '.docx', '.txt']
        
        # RFC 2119 keywords for requirement levels
        self.level_patterns = {
            PolicyLevel.MUST: r'\b(MUST|REQUIRED|SHALL)\b',
            PolicyLevel.MUST_NOT: r'\b(MUST NOT|SHALL NOT)\b',
            PolicyLevel.SHOULD: r'\b(SHOULD|RECOMMENDED)\b',
            PolicyLevel.SHOULD_NOT: r'\b(SHOULD NOT|NOT RECOMMENDED)\b',
            PolicyLevel.MAY: r'\b(MAY|OPTIONAL)\b',
        }
        
        # Category keywords for classification
        self.category_keywords = {
            PolicyCategory.SECURITY: ['security', 'authentication', 'authorization', 'encryption', 'vulnerability', 'threat', 'CVE', 'OWASP'],
            PolicyCategory.QUALITY: ['quality', 'maintainability', 'readability', 'complexity', 'duplication', 'smell'],
            PolicyCategory.PERFORMANCE: ['performance', 'latency', 'throughput', 'memory', 'CPU', 'optimization', 'speed'],
            PolicyCategory.ARCHITECTURE: ['architecture', 'design', 'pattern', 'component', 'module', 'layer', 'coupling', 'cohesion'],
            PolicyCategory.TESTING: ['test', 'coverage', 'unit test', 'integration test', 'E2E', 'assertion'],
            PolicyCategory.DOCUMENTATION: ['documentation', 'comment', 'docstring', 'README', 'API doc', 'javadoc'],
            PolicyCategory.DEPLOYMENT: ['deployment', 'CI/CD', 'pipeline', 'release', 'container', 'Docker', 'Kubernetes'],
            PolicyCategory.MAINTENANCE: ['maintenance', 'dependency', 'update', 'upgrade', 'version', 'deprecation'],
            PolicyCategory.COMPLIANCE: ['compliance', 'regulation', 'standard', 'GDPR', 'HIPAA', 'SOC2', 'PCI-DSS'],
        }
        
        # Threshold patterns (e.g., "coverage MUST be > 80%")
        self.threshold_pattern = r'([\d.]+)\s*(%|ms|MB|KB|seconds?|minutes?)'
    
    def analyze_file(self, file_path: str) -> PolicyDocument:
        """
        Analyze a policy document file.
        
        Args:
            file_path: Path to policy document
        
        Returns:
            PolicyDocument with extracted rules
        
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format not supported
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Policy file not found: {file_path}")
        
        if path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported format: {path.suffix}. Supported: {self.supported_formats}")
        
        # Calculate file hash
        file_hash = self._calculate_file_hash(path)
        
        # Parse document based on format
        if path.suffix.lower() == '.pdf':
            content = self._parse_pdf(path)
        elif path.suffix.lower() == '.md':
            content = self._parse_markdown(path)
        elif path.suffix.lower() == '.docx':
            content = self._parse_docx(path)
        else:  # .txt
            content = self._parse_text(path)
        
        # Extract metadata
        metadata = self._extract_metadata(content)
        
        # Extract policy rules
        rules = self._extract_rules(content)
        
        # Build policy document
        doc = PolicyDocument(
            file_path=str(path),
            file_hash=file_hash,
            format=path.suffix.lower()[1:],  # Remove leading dot
            title=metadata.get('title'),
            version=metadata.get('version'),
            date=metadata.get('date'),
            author=metadata.get('author'),
            scope=metadata.get('scope'),
            rules=rules,
            metadata=metadata
        )
        
        return doc
    
    def _calculate_file_hash(self, path: Path) -> str:
        """Calculate SHA256 hash of file"""
        sha256 = hashlib.sha256()
        with open(path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                sha256.update(chunk)
        return sha256.hexdigest()
    
    def _parse_pdf(self, path: Path) -> str:
        """Parse PDF document"""
        try:
            import PyPDF2
            
            content = []
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
            
            return '\n'.join(content)
        
        except ImportError:
            # Fallback: try pdfplumber
            try:
                import pdfplumber
                
                content = []
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            content.append(text)
                
                return '\n'.join(content)
            
            except ImportError:
                raise ImportError("PDF parsing requires PyPDF2 or pdfplumber. Install with: pip install PyPDF2 pdfplumber")
    
    def _parse_markdown(self, path: Path) -> str:
        """Parse Markdown document"""
        try:
            import mistune
            
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Use mistune to parse markdown, but we want plain text
            # So we'll just use the raw content for now
            return content
        
        except ImportError:
            # Fallback to plain text parsing
            with open(path, 'r', encoding='utf-8') as f:
                return f.read()
    
    def _parse_docx(self, path: Path) -> str:
        """Parse DOCX document"""
        try:
            from docx import Document
            
            doc = Document(path)
            content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            return '\n'.join(content)
        
        except ImportError:
            raise ImportError("DOCX parsing requires python-docx. Install with: pip install python-docx")
    
    def _parse_text(self, path: Path) -> str:
        """Parse plain text document"""
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _extract_metadata(self, content: str) -> Dict[str, Any]:
        """Extract document metadata"""
        metadata = {}
        
        # Extract title (first heading or line with "Title:")
        title_match = re.search(r'(?:^|\n)#\s+(.+)|Title:\s*(.+)', content, re.IGNORECASE)
        if title_match:
            metadata['title'] = (title_match.group(1) or title_match.group(2)).strip()
        
        # Extract version
        version_match = re.search(r'Version:\s*([\d.]+)', content, re.IGNORECASE)
        if version_match:
            metadata['version'] = version_match.group(1)
        
        # Extract date
        date_match = re.search(r'Date:\s*(\d{4}-\d{2}-\d{2})', content, re.IGNORECASE)
        if date_match:
            try:
                metadata['date'] = datetime.strptime(date_match.group(1), '%Y-%m-%d')
            except ValueError:
                pass
        
        # Extract author
        author_match = re.search(r'Author:\s*(.+)', content, re.IGNORECASE)
        if author_match:
            metadata['author'] = author_match.group(1).strip()
        
        # Extract scope
        scope_match = re.search(r'Scope:\s*(.+)', content, re.IGNORECASE)
        if scope_match:
            metadata['scope'] = scope_match.group(1).strip()
        
        return metadata
    
    def _extract_rules(self, content: str) -> List[PolicyRule]:
        """Extract policy rules from content"""
        rules = []
        
        # Split content into lines
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            # Check if line contains a policy level keyword
            level = self._detect_level(line)
            if not level:
                continue
            
            # Extract rule text
            rule_text = line.strip()
            
            # Generate rule ID
            rule_id = f"RULE-{len(rules) + 1:03d}"
            
            # Detect category
            category = self._detect_category(rule_text)
            
            # Extract keywords
            keywords = self._extract_keywords(rule_text)
            
            # Extract threshold if present
            threshold, unit = self._extract_threshold(rule_text)
            
            # Look for rationale (next line might start with "Rationale:" or "Because:")
            rationale = None
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if re.match(r'(Rationale|Because|Reason):', next_line, re.IGNORECASE):
                    rationale = re.sub(r'(Rationale|Because|Reason):\s*', '', next_line, flags=re.IGNORECASE)
            
            # Create policy rule
            rule = PolicyRule(
                id=rule_id,
                text=rule_text,
                level=level,
                category=category,
                keywords=keywords,
                threshold=threshold,
                unit=unit,
                rationale=rationale,
                source_line=i + 1
            )
            
            rules.append(rule)
        
        return rules
    
    def _detect_level(self, text: str) -> Optional[PolicyLevel]:
        """Detect policy level from text"""
        # Check MUST NOT before MUST (order matters)
        if re.search(self.level_patterns[PolicyLevel.MUST_NOT], text, re.IGNORECASE):
            return PolicyLevel.MUST_NOT
        
        # Check SHOULD NOT before SHOULD
        if re.search(self.level_patterns[PolicyLevel.SHOULD_NOT], text, re.IGNORECASE):
            return PolicyLevel.SHOULD_NOT
        
        # Check other levels
        for level, pattern in self.level_patterns.items():
            if level in [PolicyLevel.MUST_NOT, PolicyLevel.SHOULD_NOT]:
                continue  # Already checked
            if re.search(pattern, text, re.IGNORECASE):
                return level
        
        return None
    
    def _detect_category(self, text: str) -> PolicyCategory:
        """Detect policy category from text"""
        text_lower = text.lower()
        
        # Count keyword matches per category
        category_scores = {}
        for category, keywords in self.category_keywords.items():
            score = sum(1 for keyword in keywords if keyword.lower() in text_lower)
            if score > 0:
                category_scores[category] = score
        
        # Return category with highest score
        if category_scores:
            return max(category_scores.items(), key=lambda x: x[1])[0]
        
        return PolicyCategory.GENERAL
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract important keywords from text"""
        # Simple extraction: words in ALL CAPS or quoted
        keywords = []
        
        # ALL CAPS words (but not RFC 2119 keywords)
        caps_words = re.findall(r'\b[A-Z]{2,}\b', text)
        rfc_keywords = {'MUST', 'SHOULD', 'MAY', 'NOT', 'REQUIRED', 'SHALL', 'RECOMMENDED', 'OPTIONAL'}
        keywords.extend([w for w in caps_words if w not in rfc_keywords])
        
        # Quoted phrases
        quoted = re.findall(r'"([^"]+)"', text)
        keywords.extend(quoted)
        
        return list(set(keywords))  # Remove duplicates
    
    def _extract_threshold(self, text: str) -> Tuple[Optional[float], Optional[str]]:
        """Extract numeric threshold and unit"""
        match = re.search(self.threshold_pattern, text)
        if match:
            try:
                value = float(match.group(1))
                unit = match.group(2)
                return value, unit
            except ValueError:
                pass
        
        return None, None
    
    def get_rules_by_level(self, doc: PolicyDocument, level: PolicyLevel) -> List[PolicyRule]:
        """Filter rules by level"""
        return [rule for rule in doc.rules if rule.level == level]
    
    def get_rules_by_category(self, doc: PolicyDocument, category: PolicyCategory) -> List[PolicyRule]:
        """Filter rules by category"""
        return [rule for rule in doc.rules if rule.category == category]
    
    def get_critical_rules(self, doc: PolicyDocument) -> List[PolicyRule]:
        """Get critical rules (MUST and MUST NOT)"""
        return [rule for rule in doc.rules if rule.level in [PolicyLevel.MUST, PolicyLevel.MUST_NOT]]


def main():
    """Test policy analyzer"""
    analyzer = PolicyAnalyzer()
    
    # Create sample policy file
    sample_policy = """# Code Quality Policy
Version: 1.0
Date: 2025-11-26
Author: CORTEX Team
Scope: All Python projects

## Testing Requirements

- Test coverage MUST be greater than 80%.
Rationale: High coverage ensures code reliability.

- Unit tests SHOULD run in less than 5 seconds.

- Integration tests MAY be skipped in development environments.

## Security Requirements

- All user input MUST be validated and sanitized.

- Passwords MUST NOT be stored in plain text.

- Authentication tokens SHOULD expire after 24 hours.

## Performance Requirements

- API response time MUST be under 200ms for 95th percentile.

- Memory usage SHOULD NOT exceed 512MB per process.
"""
    
    # Write sample to temp file
    import tempfile
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(sample_policy)
        temp_path = f.name
    
    try:
        # Analyze policy
        doc = analyzer.analyze_file(temp_path)
        
        print(f"Policy Document: {doc.title}")
        print(f"Version: {doc.version}")
        print(f"Rules found: {len(doc.rules)}\n")
        
        # Show critical rules
        critical = analyzer.get_critical_rules(doc)
        print(f"Critical Rules (MUST/MUST NOT): {len(critical)}")
        for rule in critical:
            print(f"  - [{rule.level.value}] {rule.text}")
            if rule.threshold:
                print(f"    Threshold: {rule.threshold} {rule.unit}")
        
        # Show by category
        print("\nRules by Category:")
        for category in PolicyCategory:
            rules = analyzer.get_rules_by_category(doc, category)
            if rules:
                print(f"  {category.value}: {len(rules)} rules")
    
    finally:
        # Clean up
        os.unlink(temp_path)


if __name__ == "__main__":
    main()
